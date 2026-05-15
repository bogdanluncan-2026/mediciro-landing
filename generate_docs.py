"""
Generare documente Word pentru MediciRO - Ghiduri utilizatori
Rulare: python generate_docs.py
Output: public/docs/*.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "public/docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ── Stiluri comune ────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)  # albastru inchis
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x37, 0x51, 0x8c)
    return p

def add_paragraph(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def add_step(doc, number, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, indent=False):
    p = doc.add_paragraph(style='List Bullet')
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_screenshot_placeholder(doc, label="Screenshot"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"[ SCREENSHOT: {label} ]")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    # Adaugam un dreptunghi vizual simplu prin text border
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Linie de demarcatie
    doc.add_paragraph("─" * 80).runs[0].font.color.rgb = RGBColor(0xc0, 0xca, 0xd8)
    return p

def add_spacer(doc):
    doc.add_paragraph()

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("📌 Notă: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.italic = True
    return p

def add_tip(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("💡 Sfat: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    return p

def add_title_page(doc, title, subtitle, role_color):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MediciRO")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.color.rgb = role_color

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(subtitle)
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph()
    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("mediciro.ro  ·  contact@mediciro.ro")
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1: ADMINISTRATOR CLINICĂ
# ══════════════════════════════════════════════════════════════════════════════

def create_admin_doc():
    doc = Document()
    add_title_page(doc,
        "Ghid Administrator Clinică",
        "Instrucțiuni complete de utilizare a platformei MediciRO",
        RGBColor(0x1e, 0x3a, 0x8a))

    # ── 1. Înregistrare ──────────────────────────────────────────────────────
    set_heading(doc, "1. Cum înregistrezi clinica ta", 1)
    add_paragraph(doc, "Înregistrarea se face o singură dată și durează aproximativ 5 minute. Ai nevoie de o adresă de email validă și datele de bază ale clinicii.")
    add_spacer(doc)

    set_heading(doc, "1.1. Accesează pagina de înregistrare", 2)
    add_step(doc, 1, 'Deschide browserul și mergi la adresa: app.mediciro.ro')
    add_step(doc, 2, 'Apasă butonul "Înregistrează clinica" din centrul paginii.')
    add_step(doc, 3, 'Vei ajunge pe pagina de creare cont.')
    add_screenshot_placeholder(doc, "Pagina principală app.mediciro.ro cu butonul Înregistrează clinica")

    set_heading(doc, "1.2. Completează formularul", 2)
    add_paragraph(doc, "Completează câmpurile cerute:")
    add_bullet(doc, "Numele clinicii — exact cum apare pe ușa cabinetului")
    add_bullet(doc, "CUI — codul unic de înregistrare al firmei")
    add_bullet(doc, "Adresa și orașul")
    add_bullet(doc, "Numărul de telefon al clinicii")
    add_bullet(doc, "Adresa de email — aceasta va fi și adresa de autentificare")
    add_bullet(doc, "Parola — alege una sigură, cu litere mari, mici și cifre")
    add_screenshot_placeholder(doc, "Formularul de înregistrare completat")

    add_step(doc, 1, 'Bifează că ești de acord cu Termenii și Condițiile.')
    add_step(doc, 2, 'Apasă butonul "Creează cont".')
    add_note(doc, "Toate câmpurile marcate cu * sunt obligatorii.")

    set_heading(doc, "1.3. Verificarea contului și aprobarea", 2)
    add_paragraph(doc, "După trimiterea formularului, contul tău intră în revizie la echipa MediciRO.")
    add_step(doc, 1, "Vei primi un email de confirmare că cererea a fost primită.")
    add_step(doc, 2, "Echipa MediciRO verifică datele clinicii în 1-2 zile lucrătoare.")
    add_step(doc, 3, "Când contul este aprobat, primești un email cu confirmarea activării.")
    add_step(doc, 4, "Din acel moment te poți autentifica și începe să folosești platforma.")
    add_note(doc, "Dacă nu primești niciun email în 24 de ore, verifică și dosarul Spam.")

    add_spacer(doc)

    # ── 2. Autentificare ─────────────────────────────────────────────────────
    set_heading(doc, "2. Cum te autentifici", 1)
    add_step(doc, 1, 'Mergi la app.mediciro.ro și apasă "Autentificare" → "Clinică".')
    add_step(doc, 2, "Introdu adresa de email și parola alese la înregistrare.")
    add_step(doc, 3, 'Apasă "Intră în cont".')
    add_screenshot_placeholder(doc, "Pagina de autentificare clinică")
    add_tip(doc, "Dacă ai uitat parola, apasă pe legătura 'Am uitat parola' și urmează instrucțiunile primite pe email.")

    add_spacer(doc)

    # ── 3. Ecranul principal (Dashboard) ─────────────────────────────────────
    set_heading(doc, "3. Ecranul principal — ce găsești", 1)
    add_paragraph(doc, "Imediat după autentificare ajungi pe ecranul principal, numit Dashboard. Aici ai o imagine de ansamblu asupra activității clinicii.")
    add_screenshot_placeholder(doc, "Dashboard-ul administratorului cu cardurile de statistici")

    add_paragraph(doc, "Ce poți vedea pe ecranul principal:")
    add_bullet(doc, "Total programări astăzi — câte consultații sunt programate în ziua curentă")
    add_bullet(doc, "Programări în așteptare — programări care nu au fost încă confirmate")
    add_bullet(doc, "Total doctori — câți medici sunt înregistrați în clinică")
    add_bullet(doc, "Total pacienți — câți pacienți unici sunt în baza de date")
    add_paragraph(doc, "Apasă pe oricare card pentru a vedea lista detaliată.")

    add_spacer(doc)

    # ── 4. Gestionare doctori ─────────────────────────────────────────────────
    set_heading(doc, "4. Gestionarea doctorilor", 1)
    add_paragraph(doc, 'Mergi la secțiunea "Doctori" din meniul din stânga.')

    set_heading(doc, "4.1. Cum adaugi un doctor nou", 2)
    add_step(doc, 1, 'Apasă butonul "+ Adaugă doctor" din colțul din dreapta sus.')
    add_step(doc, 2, "Introdu adresa de email a doctorului. Platforma verifică automat dacă doctorul există deja în sistem.")
    add_step(doc, 3, "Completează datele:")
    add_bullet(doc, "Numele și prenumele", indent=True)
    add_bullet(doc, "Specialitatea medicală", indent=True)
    add_bullet(doc, "Codul CMR (numărul de pe avizul de liberă practică)", indent=True)
    add_bullet(doc, "Tariful pentru o consultație (în RON)", indent=True)
    add_bullet(doc, "O scurtă descriere a doctorului (opțional)", indent=True)
    add_step(doc, 4, 'Apasă "Salvează".')
    add_step(doc, 5, "Doctorul primește automat un email de invitație. Trebuie să acceseze linkul din email pentru a-și activa contul.")
    add_screenshot_placeholder(doc, "Formularul de adăugare doctor")
    add_note(doc, "Dacă doctorul lucrează la mai multe clinici și are deja cont pe MediciRO, se adaugă automat la clinica ta fără să creeze un cont nou.")

    set_heading(doc, "4.2. Cum editezi un doctor", 2)
    add_step(doc, 1, "Găsește doctorul în lista de doctori.")
    add_step(doc, 2, 'Apasă pictograma de editare (creion) din dreptul numelui.')
    add_step(doc, 3, "Modifică informațiile dorite.")
    add_step(doc, 4, 'Apasă "Salvează".')

    set_heading(doc, "4.3. Cum dezactivezi un doctor", 2)
    add_paragraph(doc, "Dacă un doctor nu mai lucrează în clinică, îl poți dezactiva. Programările existente nu se șterg.")
    add_step(doc, 1, "Găsește doctorul în listă.")
    add_step(doc, 2, 'Apasă pictograma de dezactivare sau butonul "Dezactivează".')
    add_step(doc, 3, "Confirmă acțiunea.")

    add_spacer(doc)

    # ── 5. Gestionare programări ──────────────────────────────────────────────
    set_heading(doc, "5. Gestionarea programărilor", 1)
    add_paragraph(doc, 'Mergi la secțiunea "Programări" din meniu. Programările sunt împărțite în două grupe: Următoarele programări (viitoare) și Trecute.')
    add_screenshot_placeholder(doc, "Lista de programări cu filtrele disponibile")

    set_heading(doc, "5.1. Cum filtrezi programările după status", 2)
    add_paragraph(doc, "Folosește butoanele din dreapta sus pentru a filtra după status:")
    add_bullet(doc, "Toate — afișează toate programările, grupate în Viitoare și Trecute")
    add_bullet(doc, "În așteptare — programări primite online, care nu au fost încă confirmate (marcate cu galben)")
    add_bullet(doc, "Confirmate — programări acceptate de clinică")
    add_bullet(doc, "Finalizate — consultații care au avut loc")
    add_bullet(doc, "Anulate — programări anulate de clinică sau de pacient")
    add_screenshot_placeholder(doc, "Butoanele de filtrare după status")

    set_heading(doc, "5.2. Cum filtrezi după doctor", 2)
    add_paragraph(doc, "Dacă clinica ta are mai mulți medici, poți vedea programările unui singur doctor.")
    add_step(doc, 1, 'Deasupra listei, găsești meniul derulant "Doctor:".')
    add_step(doc, 2, 'Apasă pe meniu și alege doctorul dorit din listă.')
    add_step(doc, 3, "Lista se actualizează imediat — vei vedea doar programările acelui medic.")
    add_step(doc, 4, 'Selectează "Toți doctorii" pentru a reveni la lista completă.')
    add_note(doc, "Filtrul după doctor apare doar dacă clinica are cel puțin 2 medici înregistrați.")
    add_screenshot_placeholder(doc, "Meniul derulant pentru filtrare după doctor")

    set_heading(doc, "5.3. Cum confirmi o programare", 2)
    add_paragraph(doc, "Când un pacient se programează online, programarea apare cu statusul 'În așteptare'. Trebuie să o confirmi manual.")
    add_step(doc, 1, "Identifică programările marcate cu galben (În așteptare).")
    add_step(doc, 2, 'Apasă butonul verde "Confirmă" din dreptul programării.')
    add_step(doc, 3, "Statusul se schimbă în 'Confirmată'.")
    add_step(doc, 4, "Pacientul primește automat o notificare prin SMS și email.")
    add_screenshot_placeholder(doc, "Butonul Confirmă pe o programare în așteptare")
    add_tip(doc, "Poți confirma rapid toate programările din ziua curentă filtrând după statusul 'În așteptare'.")

    set_heading(doc, "5.4. Cum anulezi o programare", 2)
    add_step(doc, 1, "Găsește programarea pe care vrei să o anulezi.")
    add_step(doc, 2, 'Apasă butonul roșu "Anulează".')
    add_step(doc, 3, "Se deschide o fereastră — introdu motivul anulării (obligatoriu).")
    add_step(doc, 4, 'Apasă "Confirmă anularea".')
    add_step(doc, 5, "Pacientul primește automat o notificare cu motivul anulării.")
    add_note(doc, "O programare anulată nu poate fi reactivată. Dacă pacientul dorește o altă dată, trebuie să creeze o programare nouă.")

    add_spacer(doc)

    # ── 6. Calendar și programări manuale ─────────────────────────────────────
    set_heading(doc, "6. Calendarul și programările manuale", 1)
    add_paragraph(doc, 'Secțiunea "Calendar" îți oferă o vedere de ansamblu vizuală a tuturor programărilor, organizate pe zile. Poți de asemenea crea programări manuale direct din calendar — util când un pacient sună sau vine la recepție.')
    add_screenshot_placeholder(doc, "Calendarul cu programările pe zile")

    set_heading(doc, "6.1. Navigarea în calendar", 2)
    add_step(doc, 1, 'Apasă pe "Calendar" în meniul din stânga.')
    add_step(doc, 2, "Folosește săgețile pentru a naviga înainte sau înapoi.")
    add_step(doc, 3, "Apasă pe o zi pentru a vedea programările din acea zi.")
    add_paragraph(doc, "Programările sunt codificate cu culori după status:")
    add_bullet(doc, "Galben — În așteptare")
    add_bullet(doc, "Albastru — Confirmate")
    add_bullet(doc, "Verde — Finalizate")
    add_bullet(doc, "Roșu — Anulate")

    set_heading(doc, "6.2. Cum creezi o programare manuală", 2)
    add_paragraph(doc, "Folosești această funcție când un pacient sună sau vine personal la recepție.")
    add_step(doc, 1, 'Apasă butonul "+ Programare nouă" sau apasă pe un slot liber în calendar.')
    add_step(doc, 2, "Selectează doctorul.")
    add_step(doc, 3, "Alege data și intervalul orar disponibil.")
    add_step(doc, 4, "Caută pacientul după nume sau telefon în câmpul de căutare.")
    add_step(doc, 5, "Dacă pacientul nu există în sistem, apasă 'Pacient nou' și completează datele.")
    add_step(doc, 6, "Alege tipul consultației: Față în față sau Teleconsultație.")
    add_step(doc, 7, "Adaugă opțional note interne (vizibile doar personalului clinicii).")
    add_step(doc, 8, 'Apasă "Confirmă programarea".')
    add_step(doc, 9, "Programarea este creată direct cu statusul Confirmată și pacientul primește notificare.")
    add_screenshot_placeholder(doc, "Formularul de creare programare manuală")
    add_note(doc, "Programările create manual de personal sunt marcate cu eticheta 'Manual' pentru a le distinge de cele făcute online de pacienți.")

    add_spacer(doc)

    # ── 7. Gestionare pacienți ────────────────────────────────────────────────
    set_heading(doc, "7. Gestionarea pacienților", 1)
    add_paragraph(doc, 'Mergi la secțiunea "Pacienți" din meniu. Aici găsești toți pacienții care au avut cel puțin o programare la clinica ta.')
    add_screenshot_placeholder(doc, "Lista de pacienți cu câmpul de căutare")

    set_heading(doc, "7.1. Cum cauți un pacient", 2)
    add_paragraph(doc, "Câmpul de căutare din dreapta sus filtrează lista în timp real.")
    add_step(doc, 1, "Apasă în câmpul de căutare.")
    add_step(doc, 2, "Scrie numele, emailul sau numărul de telefon al pacientului.")
    add_step(doc, 3, "Lista se filtrează automat pe măsură ce scrii.")
    add_step(doc, 4, 'Apasă "X" din câmpul de căutare sau șterge textul pentru a reveni la lista completă.')
    add_screenshot_placeholder(doc, "Câmpul de căutare pacienți în acțiune")

    set_heading(doc, "7.2. Informații afișate", 2)
    add_paragraph(doc, "Pentru fiecare pacient poți vedea:")
    add_bullet(doc, "Numele complet și inițialele (avatar)")
    add_bullet(doc, "Adresa de email")
    add_bullet(doc, "Numărul de telefon")
    add_bullet(doc, "Data nașterii")
    add_bullet(doc, "Asigurare privată — dacă pacientul are asigurare, apare numele asigurătorului")

    add_spacer(doc)

    # ── 8. Setările clinicii ──────────────────────────────────────────────────
    set_heading(doc, "8. Setările clinicii", 1)
    add_paragraph(doc, 'Mergi la secțiunea "Setări" din meniu. Pagina de setări are 6 tab-uri (secțiuni), fiecare pentru un aspect diferit al clinicii.')
    add_screenshot_placeholder(doc, "Pagina de setări cu cele 6 tab-uri")

    set_heading(doc, "8.1. Tab: Clinică", 2)
    add_paragraph(doc, "Informațiile de bază ale clinicii, vizibile pe facturi și în comunicările cu pacienții.")
    add_bullet(doc, "Nume clinică — denumirea oficială a clinicii")
    add_bullet(doc, "Oraș — orașul în care funcționează")
    add_bullet(doc, "Telefon — numărul principal de contact")
    add_bullet(doc, "Email clinică — adresa de email de contact")
    add_bullet(doc, "Adresă — adresa completă a clinicii")
    add_bullet(doc, "Locație pe hartă — marchează poziția exactă pe Google Maps, trăgând markerul sau căutând adresa")
    add_step(doc, 1, "Modifică câmpurile dorite.")
    add_step(doc, 2, 'Apasă "Salvează".')
    add_screenshot_placeholder(doc, "Tab Clinică — formularul cu datele de bază și harta")
    add_note(doc, "Coordonatele de pe hartă sunt folosite de pacienți pentru navigare și de platforma de booking pentru a calcula distanța.")

    set_heading(doc, "8.2. Tab: Pagina publică", 2)
    add_paragraph(doc, "Configurezi cum apare clinica ta în directorul public al platformei MediciRO, vizibil de toți pacienții.")
    add_bullet(doc, "Slug (URL clinică) — adresa web unică a clinicii, ex: app.mediciro.ro/c/clinica-mea. Alege un slug scurt și ușor de reținut, folosind doar litere mici, cifre și liniuță.")
    add_bullet(doc, "Tagline — un mesaj scurt de prezentare, maxim 1-2 propoziții (ex: 'Sănătatea ta, prioritatea noastră')")
    add_bullet(doc, "Descriere — text detaliat despre clinică: specialități, experiență, dotări, echipă")
    add_bullet(doc, "Publică pagina — comutator ON/OFF care face pagina vizibilă sau invizibilă publicului")
    add_step(doc, 1, "Completează slug-ul, tagline-ul și descrierea.")
    add_step(doc, 2, "Activează comutatorul 'Publică pagina' pentru ca pagina să fie vizibilă.")
    add_step(doc, 3, 'Apasă "Salvează".')
    add_step(doc, 4, 'Opțional: apasă "Copiază link" pentru a distribui URL-ul clinicii.')
    add_screenshot_placeholder(doc, "Tab Pagina publică — formularul și comutatorul de vizibilitate")
    add_tip(doc, "Descrierea bine scrisă ajută pacienții să aleagă clinica ta. Menționează specialitățile principale, experiența medicilor și dotările.")

    set_heading(doc, "8.3. Tab: Locații", 2)
    add_paragraph(doc, "Gestionezi punctele de lucru ale clinicii. Util dacă clinica funcționează în mai multe locuri sau clădiri.")
    add_bullet(doc, "Locația primară — creată automat la înregistrare, nu poate fi ștearsă")
    add_bullet(doc, "Locații suplimentare — adresă, telefon și email separate per punct de lucru")
    add_paragraph(doc, "Cum adaugi o locație nouă:")
    add_step(doc, 1, 'Apasă "+ Adaugă locație".')
    add_step(doc, 2, "Completează numele locației (ex: 'Punct de lucru Calea Victoriei'), telefonul, emailul, adresa și orașul.")
    add_step(doc, 3, "Opțional, marchează locația pe hartă.")
    add_step(doc, 4, 'Apasă "Salvează".')
    add_screenshot_placeholder(doc, "Tab Locații — lista locațiilor cu butonul de adăugare")
    add_note(doc, "Numărul maxim de locații depinde de planul de abonament. Dacă ai atins limita, apasă 'Upgrade plan'.")

    set_heading(doc, "8.4. Tab: Profil medic", 2)
    add_paragraph(doc, "Dacă ești și medic la clinica ta (administrator-medic), poți activa un profil de doctor care să apară în directorul de medici al platformei.")
    add_paragraph(doc, "Câmpuri necesare:")
    add_bullet(doc, "Specialitate (obligatoriu) — ex: Cardiologie, Dermatologie")
    add_bullet(doc, "Sub-specialitate (opțional) — detalii suplimentare")
    add_bullet(doc, "Cod parafă CMR (obligatoriu) — numărul de pe avizul de liberă practică")
    add_bullet(doc, "Tarif consultație (RON) — prețul pentru o consultație față în față")
    add_bullet(doc, "Tarif teleconsultație (RON) — prețul pentru o consultație video (opțional)")
    add_bullet(doc, "Bio / Descriere — scurtă prezentare a experienței și competențelor")
    add_step(doc, 1, "Completează câmpurile.")
    add_step(doc, 2, 'Apasă "Activează profil doctor".')
    add_paragraph(doc, "Dacă profilul de doctor este deja activ, apare un mesaj de confirmare și poți gestiona programul din secțiunea 'Programul meu'.")
    add_screenshot_placeholder(doc, "Tab Profil medic — formularul de activare")

    set_heading(doc, "8.5. Tab: Recepție", 2)
    add_paragraph(doc, "Setări specifice pentru personalul de la recepție (FRONTDESK): gestionarea conturilor de recepție, permisiunile acestora.")
    add_bullet(doc, "Lista conturilor de recepție active")
    add_bullet(doc, "Invitarea unui nou cont de recepție prin email")
    add_bullet(doc, "Dezactivarea unui cont de recepție")
    add_screenshot_placeholder(doc, "Tab Recepție — lista conturilor de recepție")

    set_heading(doc, "8.6. Tab: Abonament", 2)
    add_paragraph(doc, "Informații despre abonamentul activ și istoricul facturilor.")
    add_bullet(doc, "Planul activ — BASIC, PRO sau ENTERPRISE")
    add_bullet(doc, "Data reînnoirii — când expiră abonamentul curent")
    add_bullet(doc, "Funcționalități incluse — ce poți folosi cu planul tău")
    add_bullet(doc, "Istoricul facturilor — descarcă facturile lunare în format PDF")
    add_paragraph(doc, "Pentru a schimba planul sau pentru întrebări despre abonament, contactează echipa MediciRO la contact@mediciro.ro.")
    add_screenshot_placeholder(doc, "Tab Abonament — planul activ și funcționalitățile incluse")

    add_spacer(doc)

    # ── 9. Programare online pentru pacienți ──────────────────────────────────
    set_heading(doc, "9. Activarea programărilor online pentru pacienți", 1)
    add_paragraph(doc, "Poți permite pacienților să se programeze singuri, direct de pe pagina ta publică.")
    add_step(doc, 1, 'Mergi la "Setări" → tab "Pagina publică".')
    add_step(doc, 2, "Activează comutatorul 'Publică pagina'.")
    add_step(doc, 3, 'Apasă "Salvează".')
    add_note(doc, "Dacă pagina publică este dezactivată, pacienții nu pot face programări online — trebuie să te contacteze telefonic.")

    add_spacer(doc)

    # ── 10. Notificări ─────────────────────────────────────────────────────────
    set_heading(doc, "10. Notificări", 1)
    add_paragraph(doc, 'Mergi la secțiunea "Notificări" din meniu pentru a vedea toate alertele recente: programări noi, confirmări, anulări.')
    add_screenshot_placeholder(doc, "Lista de notificări")

    add_spacer(doc)

    # ── 11. Abonament ─────────────────────────────────────────────────────────
    set_heading(doc, "11. Abonamentul tău", 1)
    add_paragraph(doc, "Informațiile despre abonamentul activ se găsesc în panoul de administrare la Setări → Abonament. Dacă ai întrebări despre abonament sau vrei să îl schimbi, contactează-ne la contact@mediciro.ro.")

    doc.save(f"{OUTPUT_DIR}/Ghid_Administrator_Clinica.docx")
    print("✓ Ghid_Administrator_Clinica.docx")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2: RECEPȚIE / FRONTDESK
# ══════════════════════════════════════════════════════════════════════════════

def create_frontdesk_doc():
    doc = Document()
    add_title_page(doc,
        "Ghid Personal Recepție",
        "Instrucțiuni de utilizare a platformei MediciRO pentru recepționiști",
        RGBColor(0x05, 0x7a, 0x55))

    set_heading(doc, "Introducere", 1)
    add_paragraph(doc, "Ca personal de recepție, folosești MediciRO în fiecare zi pentru a gestiona programările pacienților. Acest ghid îți explică pas cu pas tot ce ai nevoie să știi.")

    add_spacer(doc)

    # ── 1. Autentificare ──────────────────────────────────────────────────────
    set_heading(doc, "1. Cum te autentifici", 1)
    add_step(doc, 1, "Deschide browserul și mergi la app.mediciro.ro")
    add_step(doc, 2, 'Apasă "Autentificare" → "Clinică".')
    add_step(doc, 3, "Introdu adresa de email și parola primite de la administratorul clinicii.")
    add_step(doc, 4, 'Apasă "Intră în cont".')
    add_screenshot_placeholder(doc, "Pagina de autentificare")
    add_tip(doc, "Poți salva adresa app.mediciro.ro ca marcaj în browser pentru a o accesa mai repede.")

    add_spacer(doc)

    # ── 2. Ecranul principal ──────────────────────────────────────────────────
    set_heading(doc, "2. Ecranul principal — ce găsești", 1)
    add_paragraph(doc, "După autentificare ajungi pe ecranul principal. Îți arată o privire de ansamblu a zilei curente:")
    add_bullet(doc, "Câte programări sunt astăzi")
    add_bullet(doc, "Câte sunt în așteptare de confirmare")
    add_bullet(doc, "Ultimele activități")
    add_screenshot_placeholder(doc, "Ecranul principal al recepționistului")

    add_spacer(doc)

    # ── 3. Programările ───────────────────────────────────────────────────────
    set_heading(doc, "3. Gestionarea programărilor", 1)
    add_paragraph(doc, 'Apasă pe "Programări" în meniul din stânga.')
    add_screenshot_placeholder(doc, "Lista tuturor programărilor")

    set_heading(doc, "3.1. Cum găsești o programare", 2)
    add_paragraph(doc, "Poți filtra lista de programări în mai multe feluri:")
    add_bullet(doc, "Alege data din calendar")
    add_bullet(doc, "Selectează un anumit doctor")
    add_bullet(doc, "Filtrează după status: Toate / În așteptare / Confirmate / Finalizate / Anulate")

    set_heading(doc, "3.2. Cum creezi o programare pentru un pacient", 2)
    add_paragraph(doc, "Folosești această funcție când un pacient sună sau vine la recepție și vrea să se programeze.")
    add_step(doc, 1, 'Apasă butonul "+ Programare nouă" din colțul din dreapta sus.')
    add_step(doc, 2, "Selectează doctorul la care se programează.")
    add_step(doc, 3, "Alege data dorită din calendar.")
    add_step(doc, 4, "Alege intervalul orar disponibil.")
    add_step(doc, 5, "Caută pacientul după nume sau număr de telefon.")
    add_step(doc, 6, "Dacă pacientul nu există, apasă 'Pacient nou' și completează datele lui.")
    add_step(doc, 7, "Alege tipul consultației: față în față sau teleconsultație (video).")
    add_step(doc, 8, 'Apasă "Confirmă programarea".')
    add_step(doc, 9, "Pacientul primește automat o notificare prin SMS și email.")
    add_screenshot_placeholder(doc, "Formularul de creare programare manuală")

    set_heading(doc, "3.3. Cum confirmi o programare venită online", 2)
    add_paragraph(doc, "Când un pacient se programează singur de pe site, programarea apare cu statusul 'În așteptare'. Trebuie să o confirmi.")
    add_step(doc, 1, "Găsește programarea cu statusul 'În așteptare' (marcată cu galben).")
    add_step(doc, 2, 'Apasă butonul "Confirmă".')
    add_step(doc, 3, "Pacientul primește automat un SMS și email de confirmare.")
    add_screenshot_placeholder(doc, "Butonul de confirmare pe o programare în așteptare")

    set_heading(doc, "3.4. Cum anulezi o programare", 2)
    add_step(doc, 1, "Găsește programarea în listă.")
    add_step(doc, 2, 'Apasă butonul "Anulează".')
    add_step(doc, 3, "Adaugă un motiv dacă este necesar.")
    add_step(doc, 4, "Pacientul este notificat automat.")
    add_note(doc, "O programare anulată nu se poate reactiva. Dacă pacientul dorește o altă dată, creează o programare nouă.")

    add_spacer(doc)

    # ── 4. Pacienții ──────────────────────────────────────────────────────────
    set_heading(doc, "4. Gestionarea pacienților", 1)
    add_paragraph(doc, 'Apasă pe "Pacienți" în meniu.')
    add_screenshot_placeholder(doc, "Lista de pacienți")

    set_heading(doc, "4.1. Cum cauți un pacient", 2)
    add_step(doc, 1, "Folosește câmpul de căutare din partea de sus.")
    add_step(doc, 2, "Poți căuta după: nume, prenume sau număr de telefon.")
    add_step(doc, 3, "Apasă pe pacient pentru a vedea istoricul programărilor lui.")

    set_heading(doc, "4.2. Cum adaugi un pacient nou", 2)
    add_step(doc, 1, 'Apasă butonul "+ Pacient nou".')
    add_step(doc, 2, "Completează:")
    add_bullet(doc, "Numele și prenumele", indent=True)
    add_bullet(doc, "Numărul de telefon", indent=True)
    add_bullet(doc, "Adresa de email (opțional)", indent=True)
    add_step(doc, 3, 'Apasă "Salvează".')
    add_tip(doc, "Numărul de telefon este important — pe el se trimit notificările prin SMS.")

    add_spacer(doc)

    # ── 5. Notificări ─────────────────────────────────────────────────────────
    set_heading(doc, "5. Notificări", 1)
    add_paragraph(doc, 'Secțiunea "Notificări" îți arată toate evenimentele recente: programări noi venite online, anulări, modificări.')
    add_bullet(doc, "Notificările necitite apar cu un punct colorat.")
    add_bullet(doc, "Apasă pe o notificare pentru a merge direct la programarea respectivă.")
    add_screenshot_placeholder(doc, "Lista de notificări cu indicatorul de necitite")

    doc.save(f"{OUTPUT_DIR}/Ghid_Receptie.docx")
    print("✓ Ghid_Receptie.docx")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 3: DOCTOR
# ══════════════════════════════════════════════════════════════════════════════

def create_doctor_doc():
    doc = Document()
    add_title_page(doc,
        "Ghid Doctor",
        "Instrucțiuni de utilizare a platformei MediciRO pentru medici",
        RGBColor(0x7c, 0x3a, 0xed))

    set_heading(doc, "Introducere", 1)
    add_paragraph(doc, "Ca medic, vei folosi MediciRO pentru a-ți vedea programul de lucru, programările cu pacienții și, dacă este cazul, pentru a face teleconsultații video. Contul tău este creat de administratorul clinicii — tu primești o invitație pe email.")

    add_spacer(doc)

    # ── 1. Activarea contului ─────────────────────────────────────────────────
    set_heading(doc, "1. Cum îți activezi contul", 1)
    add_paragraph(doc, "Administratorul clinicii te va adăuga în platformă. Vei primi un email de invitație.")
    add_step(doc, 1, "Verifică emailul — caută un mesaj de la noreply@mediciro.ro cu subiectul 'Invitație MediciRO'.")
    add_step(doc, 2, "Deschide emailul și apasă butonul 'Acceptă invitația'.")
    add_step(doc, 3, "Vei fi dus pe o pagină unde îți setezi parola personală.")
    add_step(doc, 4, "Alege o parolă și apasă 'Activează contul'.")
    add_step(doc, 5, "Contul tău este acum activ. Poți să te autentifici.")
    add_screenshot_placeholder(doc, "Emailul de invitație cu butonul Acceptă invitația")
    add_note(doc, "Linkul din email este valabil 48 de ore. Dacă a expirat, cere administratorului să retrimită invitația.")

    add_spacer(doc)

    # ── 2. Autentificare ──────────────────────────────────────────────────────
    set_heading(doc, "2. Cum te autentifici", 1)
    add_step(doc, 1, "Deschide browserul și mergi la app.mediciro.ro")
    add_step(doc, 2, 'Apasă "Autentificare" → "Clinică".')
    add_step(doc, 3, "Introdu adresa de email și parola stabilită la activarea contului.")
    add_step(doc, 4, 'Apasă "Intră în cont".')
    add_screenshot_placeholder(doc, "Pagina de autentificare")

    add_spacer(doc)

    # ── 3. Selectarea clinicii active ─────────────────────────────────────────
    set_heading(doc, "3. Selectarea clinicii active", 1)
    add_paragraph(doc, "Dacă lucrezi la mai multe clinici înregistrate pe MediciRO, vei vedea un selector de clinică în partea de sus a ecranului.")
    add_step(doc, 1, "Apasă pe numele clinicii din bara de sus.")
    add_step(doc, 2, "Alege clinica pentru care vrei să vezi programările.")
    add_paragraph(doc, "Toate informațiile afișate (programări, program) se referă la clinica selectată.")
    add_screenshot_placeholder(doc, "Selectorul de clinică din bara superioară")
    add_note(doc, "Dacă lucrezi la o singură clinică, acest selector nu apare.")

    add_spacer(doc)

    # ── 4. Ecranul principal ──────────────────────────────────────────────────
    set_heading(doc, "4. Ecranul principal — ce găsești", 1)
    add_paragraph(doc, "Ecranul principal îți arată:")
    add_bullet(doc, "Programările de astăzi — câți pacienți ai în ziua curentă")
    add_bullet(doc, "Programările în așteptare — care nu au fost încă confirmate")
    add_bullet(doc, "Un rezumat al activității recente")
    add_screenshot_placeholder(doc, "Dashboard-ul medicului")

    add_spacer(doc)

    # ── 5. Programul de lucru (sloturi orare) ─────────────────────────────────
    set_heading(doc, "5. Cum îți setezi programul de lucru", 1)
    add_paragraph(doc, 'Mergi la "Programul meu" din meniu. Aici definești în ce zile și la ce ore ești disponibil pentru programări.')
    add_screenshot_placeholder(doc, "Pagina Programul meu cu zilele săptămânii")

    set_heading(doc, "5.1. Cum adaugi un interval orar", 2)
    add_step(doc, 1, "Alege ziua săptămânii (Luni, Marți etc.).")
    add_step(doc, 2, 'Apasă "+ Adaugă interval".')
    add_step(doc, 3, "Setează ora de start și ora de final.")
    add_step(doc, 4, "Alege durata unei consultații (15, 20, 30 sau 60 de minute).")
    add_step(doc, 5, 'Apasă "Salvează".')
    add_paragraph(doc, "Platforma va calcula automat câte sloturi disponibile ai în acel interval.")
    add_screenshot_placeholder(doc, "Formularul de adăugare interval orar")
    add_tip(doc, "Poți avea mai multe intervale în aceeași zi. De exemplu: 09:00-13:00 dimineața și 15:00-18:00 după-amiaza.")

    add_spacer(doc)

    # ── 6. Programările mele ──────────────────────────────────────────────────
    set_heading(doc, "6. Programările mele", 1)
    add_paragraph(doc, 'Mergi la "Programările mele" din meniu. Vei vedea toate programările tale, filtrate pe clinica activă.')
    add_screenshot_placeholder(doc, "Lista programărilor medicului cu filtrele de dată și status")

    add_paragraph(doc, "Poți filtra după:")
    add_bullet(doc, "Dată — alege o zi anume")
    add_bullet(doc, "Status — Confirmate / În așteptare / Finalizate / Anulate")

    add_paragraph(doc, "Apasă pe o programare pentru a vedea detaliile pacientului.")

    add_spacer(doc)

    # ── 7. Teleconsultații ────────────────────────────────────────────────────
    set_heading(doc, "7. Teleconsultații (consultații video)", 1)
    add_paragraph(doc, "Teleconsultația este disponibilă pentru clinicile cu abonament PRO. Pacientul și medicul se văd printr-un apel video, fără a fi necesară prezența fizică la cabinet.")

    set_heading(doc, "7.1. Cum faci o teleconsultație", 2)
    add_step(doc, 1, "Găsește programarea de tip teleconsultație în lista ta de programări.")
    add_step(doc, 2, "La ora programată, apasă butonul 'Intră în consultație'.")
    add_step(doc, 3, "Se deschide o fereastră de apel video.")
    add_step(doc, 4, "Permite accesul la cameră și microfon dacă browserul îl cere.")
    add_step(doc, 5, "Pacientul va intra și el din linkul primit pe email/SMS.")
    add_step(doc, 6, "La finalul consultației, apasă 'Încheie consultația'.")
    add_screenshot_placeholder(doc, "Ecranul de teleconsultație video")
    add_note(doc, "Pentru o teleconsultație de calitate, asigură-te că ai o conexiune stabilă la internet și că ești într-un loc liniștit.")

    add_spacer(doc)

    # ── 8. Profilul meu ───────────────────────────────────────────────────────
    set_heading(doc, "8. Profilul meu", 1)
    add_paragraph(doc, 'Apasă pe numele tău din colțul din stânga jos pentru a accesa profilul tău.')
    add_bullet(doc, "Poți vedea clinicile la care ești asociat.")
    add_bullet(doc, "Poți schimba parola.")
    add_screenshot_placeholder(doc, "Pagina de profil a medicului")

    add_spacer(doc)

    # ── 9. Notificări ─────────────────────────────────────────────────────────
    set_heading(doc, "9. Notificări", 1)
    add_paragraph(doc, 'Secțiunea "Notificări" îți arată alertele recente: programări noi, confirmări, anulări.')
    add_bullet(doc, "Numărul de notificări necitite apare ca un punct roșu pe iconița din meniu.")
    add_screenshot_placeholder(doc, "Lista de notificări a medicului")

    doc.save(f"{OUTPUT_DIR}/Ghid_Doctor.docx")
    print("✓ Ghid_Doctor.docx")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 4: PACIENT
# ══════════════════════════════════════════════════════════════════════════════

def create_patient_doc():
    doc = Document()
    add_title_page(doc,
        "Ghid Pacient",
        "Cum te programezi online cu MediciRO — simplu, rapid, gratuit",
        RGBColor(0xd9, 0x70, 0x06))

    set_heading(doc, "Introducere", 1)
    add_paragraph(doc, "MediciRO îți permite să te programezi la medic în câteva minute, direct de pe telefon sau calculator, fără să dai telefoane. Serviciul este complet gratuit pentru pacienți.")

    add_spacer(doc)

    # ── 1. Creare cont ────────────────────────────────────────────────────────
    set_heading(doc, "1. Cum îți creezi un cont", 1)
    add_paragraph(doc, "Ai nevoie de un cont pentru a face o programare și a o vedea ulterior.")
    add_step(doc, 1, "Mergi la app.mediciro.ro sau descarcă aplicația MediciRO pe telefon.")
    add_step(doc, 2, 'Apasă "Cont nou" de la secțiunea Pentru pacienți.')
    add_step(doc, 3, "Completează:")
    add_bullet(doc, "Numele și prenumele", indent=True)
    add_bullet(doc, "Numărul de telefon", indent=True)
    add_bullet(doc, "Adresa de email", indent=True)
    add_bullet(doc, "O parolă", indent=True)
    add_step(doc, 4, 'Apasă "Creează cont".')
    add_step(doc, 5, "Contul este creat imediat. Poți să te autentifici.")
    add_screenshot_placeholder(doc, "Formularul de creare cont pacient")
    add_note(doc, "Nu uita parola — o vei folosi de fiecare dată când te autentifici. Dacă o uiți, poți folosi opțiunea 'Am uitat parola'.")

    add_spacer(doc)

    # ── 2. Autentificare ──────────────────────────────────────────────────────
    set_heading(doc, "2. Cum te autentifici", 1)
    add_step(doc, 1, "Mergi la app.mediciro.ro")
    add_step(doc, 2, 'Apasă "Autentificare" → "Pacient".')
    add_step(doc, 3, "Introdu emailul și parola.")
    add_step(doc, 4, 'Apasă "Intră în cont".')
    add_screenshot_placeholder(doc, "Pagina de autentificare pacient")

    add_spacer(doc)

    # ── 3. Ecranul principal ──────────────────────────────────────────────────
    set_heading(doc, "3. Ecranul principal", 1)
    add_paragraph(doc, "Aici găsești:")
    add_bullet(doc, "Programările tale — toate, organizate de la cea mai recentă")
    add_bullet(doc, "Butonul pentru o programare nouă")
    add_bullet(doc, "Notificările primite")
    add_screenshot_placeholder(doc, "Ecranul principal al pacientului cu lista de programări")

    add_paragraph(doc, "Poți filtra programările după status:")
    add_bullet(doc, "Toate — le vezi pe toate")
    add_bullet(doc, "În așteptare — programări care nu au fost încă confirmate de clinică")
    add_bullet(doc, "Confirmate — programări confirmate, la care urmează să mergi")
    add_bullet(doc, "Finalizate — consultații care au avut loc")
    add_bullet(doc, "Anulate — programări anulate")

    add_spacer(doc)

    # ── 4. Cum faci o programare ──────────────────────────────────────────────
    set_heading(doc, "4. Cum faci o programare nouă", 1)
    add_paragraph(doc, 'Apasă butonul "Programare nouă" sau "+" de pe ecranul principal.')
    add_screenshot_placeholder(doc, "Butonul de programare nouă pe ecranul principal")

    set_heading(doc, "Pasul 1 — Alege specialitatea sau doctorul", 2)
    add_paragraph(doc, "Ai două variante:")
    add_bullet(doc, "Caută după specialitate — apasă pe specialitatea medicală de care ai nevoie (cardiologie, dermatologie, pediatrie etc.)")
    add_bullet(doc, "Caută după doctor — dacă știi deja numele medicului, scrie-l în câmpul de căutare")
    add_screenshot_placeholder(doc, "Pasul 1 — grid cu specialitățile medicale")
    add_tip(doc, "Dacă alegi direct un doctor, sari direct la alegerea datei și orei.")

    set_heading(doc, "Pasul 2 — Alege clinica", 2)
    add_paragraph(doc, "Dacă ai ales o specialitate, vei vedea lista clinicilor din zona ta care au medici pe acea specialitate.")
    add_step(doc, 1, "Răsfoiește lista de clinici.")
    add_step(doc, 2, "Apasă pe clinica dorită.")
    add_screenshot_placeholder(doc, "Pasul 2 — lista de clinici disponibile")

    set_heading(doc, "Pasul 3 — Alege doctorul", 2)
    add_paragraph(doc, "Vei vedea medicii disponibili în clinica selectată pe specialitatea dorită.")
    add_step(doc, 1, "Apasă pe doctorul la care vrei să te programezi.")
    add_screenshot_placeholder(doc, "Pasul 3 — lista de doctori din clinică")

    set_heading(doc, "Pasul 4 — Alege data și ora", 2)
    add_step(doc, 1, "Alege data din calendar — zilele disponibile sunt marcate.")
    add_step(doc, 2, "Alege intervalul orar dorit din lista de ore disponibile.")
    add_step(doc, 3, "Alege tipul consultației:")
    add_bullet(doc, "Față în față — vii fizic la cabinet", indent=True)
    add_bullet(doc, "Video (teleconsultație) — consultație online, de acasă (dacă clinica oferă această opțiune)", indent=True)
    add_step(doc, 4, 'Apasă "Confirmă programarea".')
    add_screenshot_placeholder(doc, "Pasul 4 — calendarul cu orele disponibile")

    set_heading(doc, "Confirmare", 2)
    add_paragraph(doc, "Programarea ta a fost trimisă! Vei primi:")
    add_bullet(doc, "Un SMS de confirmare pe telefonul tău")
    add_bullet(doc, "Un email cu detaliile programării")
    add_paragraph(doc, "Clinica va confirma programarea — vei primi o notificare când este confirmată.")
    add_screenshot_placeholder(doc, "Ecranul de confirmare a programării")

    add_spacer(doc)

    # ── 5. Detaliile unei programări ──────────────────────────────────────────
    set_heading(doc, "5. Cum vezi detaliile unei programări", 1)
    add_step(doc, 1, "Apasă pe orice programare din lista de pe ecranul principal.")
    add_paragraph(doc, "Vei vedea:")
    add_bullet(doc, "Data și ora")
    add_bullet(doc, "Doctorul și clinica")
    add_bullet(doc, "Adresa clinicii cu harta (pentru consultații față în față)")
    add_bullet(doc, "Butonul pentru a intra în teleconsultație (pentru consultații video)")
    add_screenshot_placeholder(doc, "Pagina de detalii a unei programări")

    add_spacer(doc)

    # ── 6. Cum intri la o teleconsultație ─────────────────────────────────────
    set_heading(doc, "6. Cum intri la o teleconsultație", 1)
    add_paragraph(doc, "Dacă ai o programare de tip teleconsultație (consultație video):")
    add_step(doc, 1, "La ora programată, deschide aplicația sau site-ul.")
    add_step(doc, 2, "Mergi la programarea respectivă.")
    add_step(doc, 3, 'Apasă butonul "Intră în consultație".')
    add_step(doc, 4, "Se deschide o fereastră de apel video. Permite accesul la cameră și microfon.")
    add_step(doc, 5, "Medicul va intra și el în câteva momente.")
    add_step(doc, 6, "La finalul consultației, apasă 'Închide'.")
    add_screenshot_placeholder(doc, "Ecranul de teleconsultație video din perspectiva pacientului")
    add_tip(doc, "Poți face teleconsultația și de pe telefon, direct din browser — nu trebuie să instalezi nicio aplicație separată.")

    add_spacer(doc)

    # ── 7. Notificări ─────────────────────────────────────────────────────────
    set_heading(doc, "7. Notificările tale", 1)
    add_paragraph(doc, 'Apasă pe iconița de clopoțel din meniu pentru a vedea notificările.')
    add_bullet(doc, "Confirmare programare — clinica ți-a confirmat programarea")
    add_bullet(doc, "Reminder — cu 24 de ore și 1 oră înainte de consultație")
    add_bullet(doc, "Anulare — dacă clinica a anulat programarea")
    add_screenshot_placeholder(doc, "Lista de notificări a pacientului")

    add_spacer(doc)

    # ── 8. Profil ─────────────────────────────────────────────────────────────
    set_heading(doc, "8. Profilul tău", 1)
    add_paragraph(doc, 'Apasă pe pictograma de utilizator pentru a accesa profilul. Poți actualiza:')
    add_bullet(doc, "Numărul de telefon")
    add_bullet(doc, "Adresa de email")
    add_bullet(doc, "Parola")
    add_screenshot_placeholder(doc, "Pagina de profil a pacientului")

    add_spacer(doc)

    # ── Întrebări frecvente ────────────────────────────────────────────────────
    set_heading(doc, "9. Întrebări frecvente", 1)

    set_heading(doc, "Cât costă să mă programez prin MediciRO?", 2)
    add_paragraph(doc, "Nimic. Aplicația MediciRO pentru pacienți este complet gratuită.")

    set_heading(doc, "Pot anula o programare?", 2)
    add_paragraph(doc, "Da. Intră pe programare și apasă 'Anulează'. Clinica va fi notificată automat.")

    set_heading(doc, "Nu am primit SMS de confirmare. Ce fac?", 2)
    add_paragraph(doc, "Verifică că numărul de telefon din profil este corect. Dacă este corect, contactează clinica direct.")

    set_heading(doc, "Am uitat parola. Ce fac?", 2)
    add_step(doc, 1, "Pe pagina de autentificare, apasă 'Am uitat parola'.")
    add_step(doc, 2, "Introdu adresa de email.")
    add_step(doc, 3, "Vei primi un email cu instrucțiuni de resetare a parolei.")

    doc.save(f"{OUTPUT_DIR}/Ghid_Pacient.docx")
    print("✓ Ghid_Pacient.docx")


# ── Rulare ────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    create_admin_doc()
    create_frontdesk_doc()
    create_doctor_doc()
    create_patient_doc()
    print("\nToate documentele au fost generate în:", OUTPUT_DIR)
