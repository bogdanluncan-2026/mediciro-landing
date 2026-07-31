"""
Generare ghiduri de utilizare MediciRO (.docx) — unul pentru fiecare rol.

    python generate_guides.py

Fiecare ghid primeste o pagina de titlu, un CUPRINS real (camp TOC Word, cu numere de pagina si
navigare la click) si capitole pe doua niveluri. Cuprinsul se completeaza singur la deschidere:
documentul are `updateFields` pornit, deci Word actualizeaza campurile fara sa fie nevoie de F9.

Varianta PDF publicata pe site se genereaza separat, cu `generate_pdf.py`, care importa exact acelasi
continut. Fisierele .docx raman pentru cine vrea sa le editeze manual.

Inlocuieste generate_docs.py, care producea ghiduri ramase mult in urma fata de aplicatie.

Continutul reflecta aplicatia la 31 iulie 2026. Cand adaugi o functionalitate, actualizeaza si
capitolul corespunzator de aici — altfel ghidurile publicate raman in urma fara ca cineva sa observe.
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT_DIR = "public/docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

BRAND = RGBColor(0x1E, 0x3A, 0x8A)
MUTED = RGBColor(0x64, 0x74, 0x8B)

# Culoarea de accent a fiecarui ghid, ca sa se distinga pe pagina de titlu.
COLORS = {
    "admin": RGBColor(0x1D, 0x4E, 0xD8),
    "frontdesk": RGBColor(0x05, 0x96, 0x69),
    "doctor": RGBColor(0x7C, 0x3A, 0xED),
    "assistant": RGBColor(0x0E, 0x74, 0x90),
    "patient": RGBColor(0xEA, 0x58, 0x0C),
}


# ── Elemente de document ──────────────────────────────────────────────────────

def _set_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _update_fields_on_open(doc):
    """Word actualizeaza campurile (deci si cuprinsul) la deschiderea documentului."""
    settings = doc.settings.element
    tag = settings.find(qn("w:updateFields"))
    if tag is None:
        tag = OxmlElement("w:updateFields")
        settings.append(tag)
    tag.set(qn("w:val"), "true")


def title_page(doc, title, subtitle, color):
    for text, size, bold, rgb in [
        ("MediciRO", 30, True, BRAND),
        (title, 22, True, color),
        (subtitle, 12, False, MUTED),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = rgb

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("mediciro.ro  ·  contact@mediciro.ro  ·  iulie 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    doc.add_page_break()


def table_of_contents(doc):
    h = doc.add_paragraph()
    run = h.add_run("Cuprins")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = BRAND

    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    # \o "1-2" = titluri de nivel 1 si 2; \h = intrari pe care se poate da click; \z, \u = formatare web.
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    placeholder = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Deschide documentul în Word pentru a genera cuprinsul."
    placeholder.append(text)
    fld.append(placeholder)
    p._p.append(fld)
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    run = p.runs[0]
    run.font.size = Pt(17)
    run.font.color.rgb = BRAND
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    return p


def para(doc, text):
    p = doc.add_paragraph()
    p.add_run(text).font.size = Pt(11)
    return p


def steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item).font.size = Pt(11)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)


def note(doc, text, label="Notă", color=BRAND):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{label}: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = color
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.italic = True


def warn(doc, text):
    note(doc, text, label="Atenție", color=RGBColor(0xB4, 0x53, 0x09))


def render(blocks, doc):
    """Interpreteaza continutul declarat ca lista de tupluri (tip, valoare)."""
    handlers = {
        "h1": lambda v: h1(doc, v),
        "h2": lambda v: h2(doc, v),
        "p": lambda v: para(doc, v),
        "steps": lambda v: steps(doc, v),
        "bullets": lambda v: bullets(doc, v),
        "note": lambda v: note(doc, v),
        "warn": lambda v: warn(doc, v),
        "break": lambda v: doc.add_page_break(),
    }
    for kind, value in blocks:
        handlers[kind](value)


def build(filename, title, subtitle, color_key, blocks):
    doc = Document()
    _set_style(doc)
    _update_fields_on_open(doc)
    title_page(doc, title, subtitle, COLORS[color_key])
    table_of_contents(doc)
    render(blocks, doc)
    path = f"{OUTPUT_DIR}/{filename}"
    doc.save(path)
    print(f"  {path}")


# ── Fragmente comune ──────────────────────────────────────────────────────────


LOGIN_STAFF = [
    ("h1", "Autentificare"),
    ("p", "Aplicația web se deschide la adresa https://app.mediciro.ro. Aplicația mobilă (Android) se "
          "instalează din Google Play, căutând MediciRO."),
    ("h2", "Prima conectare"),
    ("steps", [
        "Deschide linkul din e-mailul de invitație primit de la clinică.",
        "Alege-ți o parolă și acceptă termenii și politica de confidențialitate.",
        "După activare, conectarea se face cu adresa de e-mail și parola alese.",
    ]),
    ("note", "Linkul de invitație este valabil o singură dată. Dacă a expirat sau l-ai pierdut, cere "
             "administratorului clinicii să retrimită invitația."),
    ("h2", "Autentificare în doi pași"),
    ("p", "Conturile de personal medical și administrativ folosesc obligatoriu autentificarea în doi "
          "pași. La prima conectare scanezi un cod QR cu o aplicație de autentificare (Google "
          "Authenticator, Microsoft Authenticator, Authy sau similar) și primești coduri de recuperare."),
    ("warn", "Păstrează codurile de recuperare într-un loc sigur. Fără telefon și fără coduri, contul nu "
             "poate fi deblocat decât scriind la contact@mediciro.ro, după verificarea identității."),
    ("h2", "Limba interfeței"),
    ("p", "Interfața este disponibilă în română, engleză și maghiară. Comutatorul RO / EN / HU se află în "
          "bara de sus, iar alegerea se reține pentru contul tău."),
]

SUPPORT = [
    ("h1", "Ajutor și contact"),
    ("p", "Pentru întrebări, sesizări sau solicitări legate de cont, scrie la contact@mediciro.ro."),
    ("p", "Ghidurile se actualizează pe măsură ce platforma evoluează. Dacă găsești o diferență între ce "
          "scrie aici și ce vezi în aplicație, anunță-ne — o corectăm."),
]


# ── Ghid: Administrator clinică ───────────────────────────────────────────────

ADMIN = [
    ("h1", "Ce poți face ca administrator de clinică"),
    ("img", ("public/screenshots/admin/01-dashboard.png", "Tabloul de bord al clinicii")),
    ("p", "Contul de administrator înregistrează clinica și o configurează. De aici adaugi medici și "
          "personal, stabilești programul, politicile de anulare și notificările, urmărești programările "
          "și administrezi abonamentul."),
    ("bullets", [
        "Datele clinicii, locația pe hartă și pagina publică de programare",
        "Medici: invitație, specialități, tarife, permisiuni",
        "Recepție și alți administratori",
        "Programări: lista completă, calendar, programări manuale",
        "Setări: remindere, politici de anulare, agenda comună",
        "Abonament, facturi și facturi suplimentare",
    ]),

    *LOGIN_STAFF,

    ("break", None),
    ("h1", "Înregistrarea și aprobarea clinicii"),
    ("steps", [
        "Completează formularul de la app.mediciro.ro/register: datele clinicii, CUI-ul și datele tale "
        "de contact.",
        "Confirmă adresa de e-mail din mesajul primit.",
        "Așteaptă aprobarea echipei MediciRO. Până atunci te poți autentifica, dar funcțiile clinicii "
        "sunt blocate.",
        "După aprobare primești un e-mail de confirmare și începe perioada de probă de 30 de zile.",
    ]),
    ("note", "Aprobarea este o verificare simplă că datele firmei sunt reale. Dacă durează neașteptat de "
             "mult, scrie la contact@mediciro.ro."),

    ("h1", "Datele clinicii și pagina publică"),
    ("p", "Din Setări completezi numele, adresa, telefonul și descrierea clinicii. Poziția pe hartă se "
          "stabilește căutând adresa sau trăgând marcajul — este ceea ce văd pacienții când aleg o "
          "clinică după distanță."),
    ("p", "Pagina publică permite pacienților să se programeze online fără să aibă cont dinainte. Este "
          "inclusă în toate planurile."),

    ("h1", "Medici și personal"),
    ("h2", "Invitarea unui medic"),
    ("img", ("public/screenshots/admin/02-doctori.png", "Lista medicilor clinicii")),
    ("steps", [
        "Mergi la Doctori și apasă Adaugă doctor.",
        "Introdu adresa de e-mail. Dacă medicul există deja în MediciRO la altă clinică, datele de "
        "identitate se completează automat și nu se pot modifica — se creează doar asocierea cu clinica ta.",
        "Completează specialitățile, tarifele și, dacă este cazul, codul de parafă.",
        "Trimite invitația. Medicul primește un e-mail cu link de activare.",
    ]),
    ("note", "Un medic poate lucra la mai multe clinici cu același cont. Orarul, tarifele și "
             "specialitățile sunt separate pentru fiecare clinică."),
    ("h2", "Codul de parafă"),
    ("p", "Codul de parafă identifică medicul și este unic în platformă. Titulatura „Dr.” apare în fața "
          "numelui doar pentru conturile care au cod de parafă; personalul medical fără parafă apare cu "
          "numele simplu. Medicul își poate completa singur codul din profil."),
    ("h2", "Permisiuni per medic"),
    ("bullets", [
        "Gestionarea propriilor tarife — medicul își poate modifica singur prețurile consultațiilor.",
        "Sesiuni instant — medicul poate porni teleconsultații instant către centrele de tratament.",
    ]),
    ("h2", "Recepție și administratori"),
    ("p", "Conturile de recepție se invită la fel ca medicii și lucrează cu programări și pacienți, fără "
          "acces la setări sau abonament. Poți invita și alți administratori; proprietarul clinicii poate "
          "transfera rolul de proprietar sau revoca accesul unui administrator."),

    ("break", None),
    ("h1", "Programări"),
    ("h2", "Lista programărilor"),
    ("img", ("public/screenshots/admin/03-lista-programari.png", "Programările clinicii, cu filtre")),
    ("p", "Lista completă a clinicii, cu filtre pe status, medic și perioadă. Fiecare programare arată "
          "pacientul, medicul, tipul (fizic sau teleconsultație) și starea plății."),
    ("h2", "Calendar și programări manuale"),
    ("img", ("public/screenshots/admin/04-calendar.png", "Calendarul, pentru programări manuale")),
    ("p", "Calendarul arată orarul pe zile și permite adăugarea unei programări direct într-un interval "
          "liber — pentru pacienții care sună sau vin la recepție."),
    ("p", "Când introduci numărul de telefon, aplicația caută automat un pacient existent cu același "
          "număr și îți propune să îl folosești, ca să nu se creeze dubluri. Numerele scrise cu spații "
          "sau liniuțe sunt acceptate."),
    ("warn", "Nu se pot crea programări în trecut, nici măcar la o oră trecută din ziua curentă."),
    ("h2", "Stări"),
    ("bullets", [
        "În așteptare — creată, dar neconfirmată încă.",
        "Confirmată — clinica a confirmat; pacientul primește notificare.",
        "Finalizată — consultația a avut loc.",
        "Anulată — de pacient sau de clinică, cu motiv.",
        "Neprezentat — pacientul nu a venit.",
    ]),
    ("p", "O programare finalizată din greșeală poate fi readusă cu un pas înapoi, cu motiv obligatoriu. "
          "Aplicația reține cine a făcut modificarea și când."),

    ("h1", "Agenda comună a clinicii"),
    ("img", ("public/screenshots/admin/05-agenda-clinicii.png", "Agenda: disponibilitatea colegilor pe o săptămână")),
    ("p", "Agenda arată, într-un singur tabel, disponibilitatea tuturor medicilor pe o săptămână. Un "
          "medic poate astfel îndruma pacientul către cel mai apropiat coleg liber și poate programa "
          "direct în orarul acestuia."),
    ("p", "Disponibilă în planul PRO. Se activează din Setări, de către administratorul clinicii, prin "
          "două comutatoare separate:"),
    ("bullets", [
        "Activarea agendei — medicii văd intervalele libere și ocupate ale colegilor.",
        "Afișarea numelor de pacienți — opțională, dezactivată implicit.",
    ]),
    ("note", "Al doilea comutator este separat intenționat. Faptul că un coleg este ocupat și cine este "
             "cu el sunt două cantități diferite de informație despre pacient. Recomandăm să îl lași "
             "oprit dacă nu ai un motiv clar, mai ales în specialitățile unde simpla existență a unei "
             "programări este o informație sensibilă."),

    ("h1", "Centre de tratament"),
    ("img", ("public/screenshots/admin/06-centre-tratament.png", "Centrele de tratament și asistenții lor")),
    ("p", "Centrele de tratament sunt locații partenere cu asistenți proprii, unde pacientul se află "
          "fizic, iar medicul se conectează video. Disponibile în planul PRO."),
    ("steps", [
        "Din Centre de Tratament, activează comutatorul Sesiuni instant.",
        "Adaugă centrele: nume, adresă, telefon.",
        "Invită asistenții pe e-mail și asignează-i centrelor.",
        "Din Doctori, acordă permisiunea de sesiuni instant medicilor care vor lansa consultații.",
    ]),
    ("p", "Pacienții de centru se înregistrează cu CNP, stocat criptat. Același pacient nu poate fi "
          "înregistrat de două ori la același centru și, la nevoie, poate fi mutat între centrele clinicii."),

    ("break", None),
    ("h1", "Notificări către pacienți"),
    ("p", "Confirmările, remindere și anulările se trimit automat. Pacienții care au aplicația "
          "instalată primesc notificare push, gratuit; ceilalți primesc SMS."),
    ("bullets", [
        "Planul BASIC include 100 de SMS-uri pe lună, planul PRO 200.",
        "SMS-urile peste pachetul inclus se facturează la 0,30 RON bucata, la sfârșitul lunii.",
        "Remindere: poți activa separat mesajul cu 24 de ore înainte și pe cel cu 2 ore înainte.",
    ]),

    ("h1", "Politici de anulare"),
    ("p", "Din Setări stabilești până cu cât timp înainte își poate anula pacientul singur programarea și "
          "în ce condiții se restituie plata unei teleconsultații. Pragurile sunt separate pentru "
          "programările fizice și pentru teleconsultații."),

    ("h1", "Abonament și facturare"),
    ("img", ("public/screenshots/admin/10-facturare.png", "Planul curent și funcționalitățile incluse")),
    ("h2", "Planuri"),
    ("bullets", [
        "BASIC — până la 3 medici, o locație, programări online, pagină publică, rapoarte de bază, "
        "100 SMS incluse.",
        "PRO — până la 15 medici, 3 locații, teleconsultații video, agenda comună, centre de tratament, "
        "rapoarte avansate, 200 SMS incluse.",
        "ENTERPRISE — medici și locații nelimitate, onboarding asistat, funcționalități la cerere.",
    ]),
    ("h2", "Plata"),
    ("p", "Abonamentul se plătește cu cardul, cu reînnoire automată, sau prin transfer bancar pe baza "
          "unei proforme. Facturile se emit automat și sunt disponibile în secțiunea Facturi."),
    ("h2", "Facturi suplimentare"),
    ("p", "Separat de abonament, la sfârșitul lunii se emite o factură pentru consumul peste pachet:"),
    ("bullets", [
        "SMS-urile trimise peste cele incluse în plan.",
        "Teleconsultațiile oferite gratuit pacientului: 5 RON/oră, contorizate la secundă, doar cât timp "
        "sunt cel puțin doi participanți în sesiune.",
        "Sesiunile instant către centrele de tratament, contorizate la fel.",
    ]),
    ("p", "Teleconsultațiile plătite de pacient cu cardul nu intră aici: acolo se reține un comision de "
          "10% din valoarea încasată."),
    ("note", "O sesiune la care cealaltă parte nu s-a conectat nu se facturează deloc — contorizarea "
             "pornește abia când sunt doi participanți în cameră."),

    *SUPPORT,
]


# ── Ghid: Recepție ────────────────────────────────────────────────────────────

FRONTDESK = [
    ("h1", "Ce poți face ca recepție"),
    ("img", ("public/screenshots/frontdesk/01-dashboard.png", "Tabloul de bord al recepției")),
    ("p", "Contul de recepție este făcut pentru lucrul zilnic cu pacienții: programări la telefon sau la "
          "ghișeu, confirmări, anulări și reprogramări. Nu are acces la setările clinicii sau la abonament."),

    *LOGIN_STAFF,

    ("break", None),
    ("h1", "Tabloul de bord"),
    ("p", "Arată programările de azi și pe cele care urmează, plus câte așteaptă confirmare. Fiecare card "
          "se deschide într-o listă din care vezi telefonul pacientului, notele și detaliile."),

    ("h1", "Programare nouă"),
    ("img", ("public/screenshots/frontdesk/02-calendar.png", "Calendarul, de unde se adaugă o programare")),
    ("steps", [
        "Deschide Calendar și programări manuale și alege medicul și ziua.",
        "Apasă pe intervalul liber dorit.",
        "Introdu numărul de telefon al pacientului și ieși din câmp.",
        "Dacă pacientul există deja, aplicația ți-l propune — confirmă, ca să nu creezi o dublură.",
        "Completează numele, eventual e-mailul, și salvează.",
    ]),
    ("p", "Numerele de telefon pot fi scrise cu spații sau liniuțe. Dacă numărul este greșit, mesajul de "
          "eroare apare chiar sub câmp."),
    ("warn", "Nu se pot face programări în trecut, nici la o oră deja trecută din ziua curentă."),
    ("h2", "Asigurare privată"),
    ("p", "Dacă pacientul are asigurare privată, bifează opțiunea și alege compania. Poți completa și "
          "CNP-ul și numărul de contract, dacă le ai."),

    ("h1", "Confirmare, anulare, reprogramare"),
    ("img", ("public/screenshots/frontdesk/03-programari.png", "Lista programărilor")),
    ("bullets", [
        "Confirmarea trimite automat notificare pacientului.",
        "Anularea cere un motiv, care ajunge în notificarea către pacient.",
        "Reprogramarea propune o dată nouă; pacientul este anunțat și sunat de recepție.",
    ]),

    ("h1", "Pacienți"),
    ("img", ("public/screenshots/frontdesk/04-pacienti.png", "Pacienții clinicii")),
    ("p", "Lista pacienților clinicii, cu căutare după nume, telefon sau CNP. De aici vezi istoricul "
          "programărilor unui pacient."),

    ("h1", "Pe telefon"),
    ("p", "Aplicația mobilă are aceleași funcții pentru recepție: tabloul de bord, programările, "
          "calendarul, pacienții și adăugarea unei programări noi, cu aceeași verificare a pacientului "
          "după numărul de telefon."),

    *SUPPORT,
]


# ── Ghid: Medic ───────────────────────────────────────────────────────────────

DOCTOR = [
    ("h1", "Ce poți face ca medic"),
    ("img", ("public/screenshots/doctor/01-dashboard.png", "Tabloul de bord al medicului")),
    ("p", "Contul de medic îți arată programările proprii și îți permite să îți stabilești orarul și "
          "concediile și să ții teleconsultații video. În funcție de setările clinicii, poți vedea și "
          "agenda colegilor sau porni sesiuni instant către centre de tratament."),

    *LOGIN_STAFF,

    ("break", None),
    ("h1", "Profilul tău"),
    ("img", ("public/screenshots/doctor/07-profil.png", "Profilul medicului, cu specialitate și cod de parafă")),
    ("p", "În profil completezi specialitatea, subspecialitatea, biografia și codul de parafă. "
          "Titulatura „Dr.” apare în fața numelui doar dacă ai cod de parafă completat."),
    ("note", "Specialitatea se alege dintr-o listă predefinită. Dacă a ta lipsește, cere-i "
             "administratorului clinicii să solicite adăugarea ei."),
    ("h2", "Mai multe clinici"),
    ("p", "Dacă lucrezi la mai multe clinici, în bara de sus ai un selector. Programările, orarul și "
          "tarifele sunt separate pe fiecare clinică, iar selecția se păstrează între pagini."),

    ("h1", "Programul de lucru"),
    ("img", ("public/screenshots/doctor/03-programul-meu.png", "Programul de lucru și concediile")),
    ("steps", [
        "Deschide Programul meu și alege clinica și locația.",
        "Adaugă intervalele de lucru pe fiecare zi a săptămânii.",
        "Stabilește durata unei consultații — din ea se generează intervalele pe care le văd pacienții.",
    ]),
    ("h2", "Concedii"),
    ("p", "Perioadele de concediu blochează generarea de intervale libere. Dacă în perioada aleasă există "
          "deja programări neîncheiate, aplicația te avertizează înainte de salvare."),
    ("h2", "Blocarea unor intervale punctuale"),
    ("p", "Pe lângă concedii, poți bloca intervale individuale direct din calendar, pentru situații "
          "punctuale."),

    ("break", None),
    ("h1", "Programările mele"),
    ("img", ("public/screenshots/doctor/02-programarile-mele.png", "Programările proprii")),
    ("p", "Lista programărilor tale, filtrată pe clinica selectată. Poți confirma, finaliza, anula sau "
          "reprograma, și poți vedea datele de contact și notele pacientului."),
    ("p", "Dacă finalizezi din greșeală o programare, o poți readuce cu un pas înapoi, indicând motivul. "
          "Aplicația reține cine a făcut modificarea și când."),

    ("h1", "Teleconsultații"),
    ("p", "La confirmarea unei teleconsultații alegi cum se plătește:"),
    ("bullets", [
        "Pacientul plătește cu cardul — achită online înainte de a se putea conecta.",
        "Scutit de plată — pacientul nu plătește nimic. Dacă teleconsultația are loc, clinica plătește "
        "către MediciRO 5 RON/oră, contorizat la secundă.",
    ]),
    ("p", "Consultația se deschide în aplicație, fără instalări. Butonul Extinde mărește fereastra video "
          "pe tot ecranul."),
    ("note", "Contorizarea pornește abia când sunteți amândoi în cameră și se oprește la închidere. O "
             "consultație la care pacientul nu s-a conectat nu se facturează."),

    ("h1", "Agenda clinicii"),
    ("img", ("public/screenshots/doctor/04-agenda-clinicii.png", "Agenda clinicii, văzută de medic")),
    ("p", "Dacă administratorul a activat-o, vezi disponibilitatea colegilor pe o săptămână și poți "
          "programa direct la un coleg liber — util când pacientul tău are nevoie de o consultație mai "
          "devreme decât ai tu loc."),
    ("bullets", [
        "Săptămâna curentă începe cu ziua de ieri, ca să ai contextul imediat.",
        "Medicii sunt ordonați după cine are primul interval liber.",
        "Poți ascunde sau afișa zilele de weekend.",
        "Pe programările tale poți deschide detaliile; la colegi vezi doar că intervalul este ocupat.",
    ]),

    ("h1", "Sesiuni instant către centre de tratament"),
    ("img", ("public/screenshots/doctor/05-sesiune-instant.png", "Lansarea unei sesiuni instant")),
    ("p", "Dacă ai primit permisiunea, poți porni o consultație video imediată către un asistent aflat "
          "lângă pacient, într-un centru partener."),
    ("steps", [
        "Deschide Sesiune instant și alege centrul.",
        "Introdu CNP-ul pacientului; dacă este deja înregistrat, numele se completează automat.",
        "Alege asistentul și lansează sesiunea — asistentul primește notificare și e-mail cu link.",
        "La final, apasă Închide sesiunea.",
    ]),
    ("warn", "Dacă asistentul nu se conectează, sesiunea se închide automat după 5 minute și nu se "
             "facturează. Vei fi avertizat cu 30 de secunde înainte."),
    ("p", "În istoric, fiecare sesiune arată motivul încheierii: finalizată, fără asistent, nimeni nu a "
          "intrat, închisă automat pentru cameră goală sau durată maximă atinsă."),

    ("h1", "Pe telefon"),
    ("p", "Aplicația mobilă acoperă programările, calendarul, orarul, concediile, agenda clinicii și "
          "lansarea sesiunilor instant. Consultația video se deschide în browserul telefonului."),

    *SUPPORT,
]


# ── Ghid: Asistent centru de tratament ────────────────────────────────────────

ASSISTANT = [
    ("h1", "Rolul de asistent de centru"),
    ("p", "Ești lângă pacient, într-un centru de tratament, iar medicul se conectează video de la "
          "distanță. Rolul tău este să primești sesiunea pornită de medic și să fii prezent în cameră "
          "împreună cu pacientul."),
    ("p", "Contul de asistent este limitat intenționat: primești și intri în sesiuni, nu le inițiezi și "
          "nu ai acces la programările sau pacienții clinicii."),

    ("h1", "Contul tău"),
    ("steps", [
        "Administratorul clinicii te invită pe e-mail și te asignează unui centru.",
        "Deschizi linkul din invitație și îți alegi o parolă.",
        "Te conectezi la https://app.mediciro.ro cu e-mailul și parola.",
    ]),
    ("p", "Ca orice cont de personal, și acesta folosește autentificarea în doi pași. La prima conectare "
          "scanezi un cod QR cu o aplicație de autentificare și primești coduri de recuperare."),
    ("warn", "Păstrează codurile de recuperare. Fără telefon și fără ele, contul nu poate fi deblocat "
             "decât scriind la contact@mediciro.ro."),
    ("p", "Poți fi asignat mai multor centre. Vei primi sesiunile pornite către oricare dintre ele."),

    ("break", None),
    ("h1", "Cum primești o sesiune"),
    ("img", ("public/screenshots/assistant/01-sesiunile-mele.png", "Sesiunile asistentului: active și istoric")),
    ("p", "Când un medic pornește o sesiune către tine primești o notificare și un e-mail cu link direct. "
          "Sesiunea apare și în pagina ta, la Sesiuni active."),
    ("steps", [
        "Deschide linkul din notificare sau apasă Intră în pagina ta.",
        "Permite accesul la cameră și microfon, dacă browserul le cere.",
        "Rămâi în cameră împreună cu pacientul pe durata consultației.",
    ]),
    ("note", "Butonul Extinde mărește fereastra video pe tot ecranul, ca să vedeți amândoi mai bine "
             "medicul."),

    ("h1", "Cât durează și ce se contorizează"),
    ("p", "Timpul se contorizează doar cât timp sunteți cel puțin doi în cameră. Cât aștepți singur, nu "
          "se contorizează nimic."),
    ("bullets", [
        "Dacă nu intră nimeni, sesiunea se închide singură după 5 minute.",
        "Dacă rămâi singur în cameră mai mult de 5 minute, sesiunea se închide la fel.",
        "Sesiunea se încheie definitiv când medicul apasă Închide sesiunea.",
    ]),

    ("h1", "Dacă pici de pe internet"),
    ("p", "Butonul Ieși te scoate din cameră, dar nu încheie sesiunea. Cât timp medicul nu a închis-o, "
          "poți reintra din pagina ta sau din același link."),
    ("warn", "Dacă ai ieșit și nu reintri, iar camera rămâne cu mai puțin de doi participanți timp de "
             "5 minute, sesiunea se închide automat și medicul va trebui să pornească una nouă."),

    ("h1", "Istoricul sesiunilor"),
    ("p", "Sub sesiunile active vezi sesiunile anterioare, cu data, durata și motivul încheierii: "
          "finalizată, fără asistent, nimeni nu a intrat sau închisă automat."),

    ("h1", "Pe telefon"),
    ("p", "Poți intra într-o sesiune și de pe telefon, deschizând linkul primit în browserul telefonului. "
          "Un ecran dedicat în aplicația mobilă este în pregătire."),

    *SUPPORT,
]


# ── Ghid: Pacient ─────────────────────────────────────────────────────────────

PATIENT = [
    ("h1", "Bine ai venit"),
    ("img", ("public/screenshots/patient/01-dashboard.png", "Programările tale")),
    ("p", "Cu MediciRO te programezi la clinica ta fără telefoane, primești reminder înainte de "
          "consultație și poți avea consultații video de acasă. Poți folosi aplicația din browser, la "
          "app.mediciro.ro, sau aplicația de Android din Google Play."),

    ("h1", "Contul tău"),
    ("h2", "Creare cont"),
    ("steps", [
        "Apasă Înregistrare și completează numele, e-mailul și parola. Poți folosi și contul Google.",
        "Numărul de telefon este opțional la înregistrare, dar îți trebuie ca să poți face programări — "
        "clinica trebuie să te poată contacta.",
        "Confirmă adresa de e-mail din mesajul primit.",
    ]),
    ("note", "Dacă ai fost deja programat de o clinică la telefon, la prima conectare aplicația îți "
             "propune programările făcute pe numărul tău, ca să le legi de cont."),
    ("h2", "Limba"),
    ("p", "Aplicația este disponibilă în română, engleză și maghiară. Alegerea ta se păstrează și se "
          "folosește și pentru e-mailurile și notificările primite."),

    ("break", None),
    ("h1", "Cum faci o programare"),
    ("img", ("public/screenshots/patient/02-programare-noua.png", "Alegerea specialității și a medicului")),
    ("steps", [
        "Apasă Programare nouă.",
        "Alege specialitatea de care ai nevoie sau caută direct medicul după nume.",
        "Alege clinica — pe hartă sau din listă, ordonată după distanță.",
        "Alege medicul, apoi ziua și ora dintre intervalele libere.",
        "Confirmă. Primești imediat confirmarea pe e-mail și în aplicație.",
    ]),
    ("p", "Programarea rămâne În așteptare până când clinica o confirmă. Vei fi anunțat când se schimbă "
          "starea."),

    ("h1", "Teleconsultații"),
    ("p", "Dacă alegi o teleconsultație, consultația are loc video, direct în aplicație — fără instalări "
          "și fără programe suplimentare."),
    ("bullets", [
        "Dacă teleconsultația se plătește, achiți online cu cardul înainte de consultație.",
        "Unele clinici pot oferi teleconsultația fără plată din partea ta.",
        "Butonul de intrare devine activ în preajma orei programate.",
    ]),
    ("note", "Asigură-te că ești într-un loc liniștit, cu internet stabil, și că permiți accesul la "
             "cameră și microfon când browserul cere."),

    ("h1", "Anulare și reprogramare"),
    ("p", "Îți poți anula singur programarea din detaliile ei, atât timp cât te încadrezi în intervalul "
          "stabilit de clinică. Fiecare clinică își stabilește propriile praguri, iar aplicația îți spune "
          "clar dacă mai poți anula și ce se întâmplă cu plata."),
    ("p", "Dacă clinica sau medicul trebuie să mute programarea, primești o propunere de dată nouă și "
          "ești contactat telefonic de recepție."),

    ("h1", "Notificări"),
    ("img", ("public/screenshots/patient/03-notificari.png", "Notificările primite")),
    ("p", "Primești confirmări, remindere înainte de consultație și înștiințări la anulare. Dacă ai "
          "aplicația instalată, ajung ca notificări pe telefon; altfel, prin SMS și e-mail."),

    ("h1", "Datele tale"),
    ("img", ("public/screenshots/patient/04-profil.png", "Profilul și datele personale")),
    ("p", "Datele medicale sunt stocate criptat, pe servere din Uniunea Europeană. Ai dreptul să îți vezi "
          "datele, să le corectezi și să îți ștergi contul, direct din profil."),
    ("p", "Consultațiile video nu sunt înregistrate de platformă."),

    *SUPPORT,
]


# ── Generare ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generare ghiduri MediciRO:")
    build("Ghid_Administrator_Clinica.docx", "Ghid Administrator Clinică",
          "Configurarea și administrarea clinicii în MediciRO", "admin", ADMIN)
    build("Ghid_Receptie.docx", "Ghid Recepție",
          "Programări, pacienți și lucrul zilnic la ghișeu", "frontdesk", FRONTDESK)
    build("Ghid_Medic.docx", "Ghid Medic",
          "Program, programări, teleconsultații și agenda clinicii", "doctor", DOCTOR)
    build("Ghid_Asistent_Centru.docx", "Ghid Asistent Centru de Tratament",
          "Primirea și desfășurarea sesiunilor instant", "assistant", ASSISTANT)
    build("Ghid_Pacient.docx", "Ghid Pacient",
          "Programări online, teleconsultații și contul tău", "patient", PATIENT)
    print("\nUrmătorul pas: python generate_pdf.py")
